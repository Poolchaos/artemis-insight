"""
Summary management routes for AI-generated document summaries.
"""

from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from motor.motor_asyncio import AsyncIOMotorDatabase
from bson import ObjectId
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from app.database import get_db
from app.models.user import UserInDB
from app.models.summary import (
    SummaryCreate,
    SummaryResponse,
    SummaryListItem,
    SummaryStatus
)
from app.models.job import JobCreate, JobType, JobStatus, JobResponse
from app.models.document import DocumentStatus
from app.middleware.auth import get_current_user
from app.services.document_service import DocumentService
from app.services.template_service import TemplateService
from app.tasks import generate_summary_task, regenerate_section_task


router = APIRouter(prefix="/api/summaries", tags=["summaries"])


@router.post("", response_model=dict, status_code=status.HTTP_202_ACCEPTED)
async def create_summary(
    document_id: str = Query(..., description="Document ID to summarize"),
    template_id: str = Query(..., description="Template ID to use for summarization"),
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Start a document summarization job.

    Creates an async Celery task to generate a summary using the specified template.
    Returns job information for status polling.

    - **document_id**: ID of the uploaded document
    - **template_id**: ID of the template to use for summarization

    Returns:
    - **job_id**: ID for tracking job status
    - **message**: Success message with status polling instructions
    """
    # Validate ObjectIds
    if not ObjectId.is_valid(document_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document_id format"
        )

    if not ObjectId.is_valid(template_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid template_id format"
        )

    # Verify document exists and belongs to user
    doc_service = DocumentService(db)
    document = await doc_service.get_document_by_user(document_id, str(current_user.id))

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify document is ready for processing
    if document.status != DocumentStatus.COMPLETED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Document status must be 'completed', current status: {document.status}"
        )

    # Verify template exists
    template_service = TemplateService(db)
    template = await template_service.get_template(template_id)

    if not template:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template not found"
        )

    if not template.is_active:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Template is not active"
        )

    # Create job record
    job_id = ObjectId()
    job_doc = {
        "_id": job_id,
        "user_id": current_user.id,
        "document_id": ObjectId(document_id),
        "template_id": ObjectId(template_id),
        "job_type": JobType.SUMMARIZE,
        "status": JobStatus.PENDING,
        "progress": 0,
        "started_at": datetime.utcnow(),
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }

    await db.jobs.insert_one(job_doc)

    # Start Celery task
    task = generate_summary_task.apply_async(
        kwargs={
            "document_id": document_id,
            "template_id": template_id,
            "user_id": str(current_user.id),
            "job_id": str(job_id)
        }
    )

    # Update job with celery_task_id
    await db.jobs.update_one(
        {"_id": job_id},
        {"$set": {"celery_task_id": task.id}}
    )

    return {
        "job_id": str(job_id),
        "celery_task_id": task.id,
        "status": JobStatus.PENDING,
        "message": f"Summarization job created. Poll GET /api/jobs/{str(job_id)} for status."
    }


@router.get("", response_model=List[SummaryListItem])
async def list_summaries(
    document_id: Optional[str] = Query(None, description="Filter by document ID"),
    template_id: Optional[str] = Query(None, description="Filter by template ID"),
    status: Optional[SummaryStatus] = Query(None, description="Filter by status"),
    skip: int = Query(0, ge=0, description="Number of records to skip"),
    limit: int = Query(20, ge=1, le=100, description="Maximum number of records to return"),
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    List summaries for the current user.

    Supports filtering by document, template, and status.
    Returns condensed summary information for list views.
    """
    # Build query
    query = {"user_id": ObjectId(current_user.id)}

    if document_id:
        if not ObjectId.is_valid(document_id):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid document_id format"
            )
        query["document_id"] = ObjectId(document_id)

    if template_id:
        if not ObjectId.is_valid(template_id):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid template_id format"
            )
        query["template_id"] = template_id

    if status:
        query["status"] = status

    # Query database
    cursor = db.summaries.find(query).sort("created_at", -1).skip(skip).limit(limit)
    summaries = await cursor.to_list(length=limit)

    # Convert to list items
    items = []
    for summary in summaries:
        section_count = len(summary.get("sections", []))
        total_word_count = sum(s.get("word_count", 0) for s in summary.get("sections", []))

        items.append(
            SummaryListItem(
                id=str(summary["_id"]),
                document_id=str(summary["document_id"]),
                template_name=summary["template_name"],
                status=summary["status"],
                section_count=section_count,
                total_word_count=total_word_count,
                started_at=summary["started_at"],
                completed_at=summary.get("completed_at")
            )
        )

    return items


@router.get("/{summary_id}", response_model=SummaryResponse)
async def get_summary(
    summary_id: str,
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Get a specific summary by ID.

    Returns complete summary with all sections, metadata, and processing information.
    """
    if not ObjectId.is_valid(summary_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary_id format"
        )

    # Query database
    summary = await db.summaries.find_one({
        "_id": ObjectId(summary_id),
        "user_id": ObjectId(current_user.id)
    })

    if not summary:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    # Convert to response model
    return SummaryResponse(
        id=str(summary["_id"]),
        document_id=str(summary["document_id"]),
        user_id=str(summary["user_id"]),
        job_id=str(summary["job_id"]) if summary.get("job_id") else None,
        template_id=summary["template_id"],
        template_name=summary["template_name"],
        status=summary["status"],
        sections=summary.get("sections", []),
        metadata=summary.get("metadata"),
        error_message=summary.get("error_message"),
        started_at=summary["started_at"],
        completed_at=summary.get("completed_at"),
        created_at=summary["created_at"],
        updated_at=summary["updated_at"]
    )


@router.delete("/{summary_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_summary(
    summary_id: str,
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Delete a summary.

    Permanently removes the summary record. Does not affect the original document.
    """
    if not ObjectId.is_valid(summary_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary_id format"
        )

    # Delete summary
    result = await db.summaries.delete_one({
        "_id": ObjectId(summary_id),
        "user_id": ObjectId(current_user.id)
    })

    if result.deleted_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    return None


@router.post("/{summary_id}/retry", response_model=dict, status_code=status.HTTP_202_ACCEPTED)
async def retry_failed_summary(
    summary_id: str,
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Retry a failed summary generation.

    Creates a new job to regenerate the summary using the same document and template.
    Only works for summaries with 'failed' status.

    Returns:
    - **job_id**: ID for tracking the new job status
    - **message**: Instructions for polling job status
    """
    if not ObjectId.is_valid(summary_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary_id format"
        )

    # Get the failed summary
    summary = await db.summaries.find_one({
        "_id": ObjectId(summary_id),
        "user_id": ObjectId(current_user.id)
    })

    if not summary:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    if summary["status"] != SummaryStatus.FAILED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Can only retry failed summaries. Current status: {summary['status']}"
        )

    document_id = str(summary["document_id"])
    template_id = summary["template_id"]

    # Verify document still exists and is ready
    doc_service = DocumentService(db)
    document = await doc_service.get_document_by_user(document_id, str(current_user.id))

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Original document not found"
        )

    if document.status != DocumentStatus.COMPLETED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Document must be 'completed' to retry. Current status: {document.status}"
        )

    # Verify template still exists
    template_service = TemplateService(db)
    template = await template_service.get_template(template_id)

    if not template:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Original template not found"
        )

    # Delete the old failed summary
    await db.summaries.delete_one({"_id": ObjectId(summary_id)})

    # Create new job record
    job_id = ObjectId()
    job_doc = {
        "_id": job_id,
        "user_id": current_user.id,
        "document_id": ObjectId(document_id),
        "template_id": ObjectId(template_id),
        "job_type": JobType.SUMMARIZE,
        "status": JobStatus.PENDING,
        "progress": 0,
        "started_at": datetime.utcnow(),
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }

    await db.jobs.insert_one(job_doc)

    # Start Celery task
    task = generate_summary_task.apply_async(
        kwargs={
            "document_id": document_id,
            "template_id": template_id,
            "user_id": str(current_user.id),
            "job_id": str(job_id)
        }
    )

    # Update job with celery_task_id
    await db.jobs.update_one(
        {"_id": job_id},
        {"$set": {"celery_task_id": task.id}}
    )

    return {
        "job_id": str(job_id),
        "celery_task_id": task.id,
        "status": JobStatus.PENDING,
        "message": f"Retry job created. Poll GET /api/jobs/{str(job_id)} for status."
    }


@router.post("/{summary_id}/regenerate-section", response_model=dict, status_code=status.HTTP_202_ACCEPTED)
async def regenerate_summary_section(
    summary_id: str,
    section_title: str = Query(..., description="Title of the section to regenerate"),
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Regenerate a specific section of a summary.

    Creates a new job to regenerate a single section using the ProcessingEngine.
    Useful for iteratively improving sections without reprocessing the entire document.

    - **summary_id**: ID of the existing summary
    - **section_title**: Exact title of the section to regenerate (e.g., "Introduction", "References")

    Returns job information for status polling.
    """
    if not ObjectId.is_valid(summary_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary_id format"
        )

    # Verify summary exists and belongs to user
    summary = await db.summaries.find_one({
        "_id": ObjectId(summary_id),
        "user_id": current_user.id
    })

    if not summary:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    # Verify summary is completed
    if summary["status"] != SummaryStatus.COMPLETED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Can only regenerate sections of completed summaries"
        )

    # Verify section exists in summary
    section_titles = [s["title"] for s in summary.get("sections", [])]
    if section_title not in section_titles:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Section '{section_title}' not found in summary. Available sections: {', '.join(section_titles)}"
        )

    # Create job record
    job_id = ObjectId()
    job_doc = {
        "_id": job_id,
        "user_id": current_user.id,
        "document_id": summary["document_id"],
        "template_id": ObjectId(summary["template_id"]),
        "summary_id": ObjectId(summary_id),
        "job_type": JobType.REGENERATE_SECTION,
        "status": JobStatus.PENDING,
        "progress": 0,
        "started_at": datetime.utcnow(),
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }

    await db.jobs.insert_one(job_doc)

    # Start Celery task
    task = regenerate_section_task.apply_async(
        kwargs={
            "summary_id": summary_id,
            "section_title": section_title,
            "user_id": str(current_user.id),
            "job_id": str(job_id)
        }
    )

    # Update job with celery_task_id
    await db.jobs.update_one(
        {"_id": job_id},
        {"$set": {"celery_task_id": task.id}}
    )

    return {
        "job_id": str(job_id),
        "celery_task_id": task.id,
        "status": JobStatus.PENDING,
        "section_title": section_title,
        "message": f"Section regeneration job created. Poll GET /api/jobs/{str(job_id)} for status."
    }


@router.get("/{summary_id}/export/pdf")
async def export_summary_pdf(
    summary_id: str,
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Export summary as PDF.
    """
    # Validate ObjectId
    try:
        summary_obj_id = ObjectId(summary_id)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary ID format"
        )

    # Get summary
    summary = await db.summaries.find_one({
        "_id": summary_obj_id,
        "user_id": ObjectId(current_user.id)
    })

    if not summary:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=30,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        spaceBefore=12
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_LEFT
    )

    # Add title
    title = Paragraph(summary.get('template_name', 'Summary'), title_style)
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))

    # Add metadata
    created_at = summary.get('created_at', datetime.utcnow()).strftime('%Y-%m-%d %H:%M')
    metadata = Paragraph(f"<b>Generated:</b> {created_at}<br/><b>Status:</b> {summary.get('status', 'unknown')}", body_style)
    elements.append(metadata)
    elements.append(Spacer(1, 0.3*inch))

    # Add sections
    sections = sorted(summary.get('sections', []), key=lambda x: x.get('order', 0))

    for section in sections:
        # Section title
        section_title = Paragraph(f"{section.get('order', '')}. {section.get('title', 'Untitled')}", heading_style)
        elements.append(section_title)

        # Section content - convert markdown to ReportLab format
        content = section.get('content', 'No content')

        # Escape HTML special characters first
        content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        # Convert markdown headers to bold text (already escaped, so use safe tags)
        import re
        content = re.sub(r'###\s*(.*?)(?=\n|$)', r'<b>\1</b>', content)
        content = re.sub(r'##\s*(.*?)(?=\n|$)', r'<b>\1</b>', content)
        content = re.sub(r'#\s*(.*?)(?=\n|$)', r'<b>\1</b>', content)

        # Convert markdown bold (pairs of **)
        content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)

        # Convert markdown italic (single *)
        content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)

        # Convert line breaks
        content = content.replace('\n\n', '<br/><br/>')
        content = content.replace('\n', '<br/>')

        content_para = Paragraph(content, body_style)
        elements.append(content_para)
        elements.append(Spacer(1, 0.2*inch))

    # Build PDF
    doc.build(elements)

    # Get PDF data
    buffer.seek(0)

    filename = f"{summary.get('template_name', 'summary').replace(' ', '_')}_{summary_id}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{summary_id}/export/docx")
async def export_summary_docx(
    summary_id: str,
    current_user: UserInDB = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """
    Export summary as Word document (DOCX).
    """
    # Validate ObjectId
    try:
        summary_obj_id = ObjectId(summary_id)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid summary ID format"
        )

    # Get summary
    summary = await db.summaries.find_one({
        "_id": summary_obj_id,
        "user_id": ObjectId(current_user.id)
    })

    if not summary:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Summary not found"
        )

    # Create Word document
    doc = Document()

    # Add title
    title = doc.add_heading(summary.get('template_name', 'Summary'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    doc.add_paragraph()
    metadata_para = doc.add_paragraph()
    metadata_para.add_run('Generated: ').bold = True
    metadata_para.add_run(summary.get('created_at', datetime.utcnow()).strftime('%Y-%m-%d %H:%M'))
    metadata_para.add_run('\nStatus: ').bold = True
    metadata_para.add_run(summary.get('status', 'unknown'))

    doc.add_paragraph()
    doc.add_paragraph()

    # Add sections
    sections = sorted(summary.get('sections', []), key=lambda x: x.get('order', 0))

    for section in sections:
        # Section heading
        section_heading = doc.add_heading(
            f"{section.get('order', '')}. {section.get('title', 'Untitled')}",
            level=1
        )

        # Section content
        content = section.get('content', 'No content')

        # Simple markdown parsing for Word
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                para = doc.add_paragraph()

                # Handle headers
                if line.startswith('###'):
                    para.style = 'Heading 3'
                    para.text = line.replace('###', '').strip()
                elif line.startswith('##'):
                    para.style = 'Heading 2'
                    para.text = line.replace('##', '').strip()
                elif line.startswith('#'):
                    para.style = 'Heading 1'
                    para.text = line.replace('#', '').strip()
                else:
                    # Handle bold and italic
                    text = line
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # Regular text, check for italic
                            italic_parts = part.split('*')
                            for j, ipart in enumerate(italic_parts):
                                if j % 2 == 0:
                                    para.add_run(ipart)
                                else:
                                    para.add_run(ipart).italic = True
                        else:
                            # Bold text
                            para.add_run(part).bold = True

        doc.add_paragraph()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{summary.get('template_name', 'summary').replace(' ', '_')}_{summary_id}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
